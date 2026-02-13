"""
Script to create mock form files for testing
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
import os

MOCK_DATA_DIR = Path(__file__).parent / "MockData"
MOCK_DATA_DIR.mkdir(exist_ok=True)

def create_handwritten_form_text(filename, production_line, status, manager, date, reason=None, workers=20):
    """Create a text file with handwritten form content (for OCR simulation)"""
    content = f"""Production Line Status Form

Production Line: {production_line}
Status: {status}
Manager: {manager}
Date: {date}
Workers: {workers}
Skills: assembly, quality_check, packaging
"""
    if reason:
        content += f"Reason: {reason}\n"
    
    # Save as text file (will be used for OCR simulation)
    with open(MOCK_DATA_DIR / filename.replace('.png', '.txt'), 'w') as f:
        f.write(content)
    print(f"Created text file for: {filename} (use any image file and name it {filename})")

def create_digital_form(filename, production_line, status, manager, date, reason=None, workers=20):
    """Create a digital DOCX form"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Production Line Status Form', 0)
    title.alignment = 1  # Center
    
    # Add content
    doc.add_paragraph(f'Production Line ID: {production_line}')
    doc.add_paragraph(f'Status: {status}')
    doc.add_paragraph(f'Manager Name: {manager}')
    doc.add_paragraph(f'Date: {date}')
    doc.add_paragraph(f'Number of Workers: {workers}')
    doc.add_paragraph('Skills Required: assembly, quality_check, packaging')
    
    if reason:
        doc.add_paragraph(f'Reason for Status: {reason}')
    
    # Save document
    doc.save(MOCK_DATA_DIR / filename)
    print(f"Created: {filename}")

# Create mock forms
print("Creating mock form files...")

# Handwritten forms (No-Go) - Create text files with content
create_handwritten_form_text(
    "form_handwritten_01.png",
    "PL-042",
    "NO-GO",
    "John Smith",
    "2025-01-15",
    "Equipment malfunction - conveyor belt stopped",
    20
)

create_handwritten_form_text(
    "form_handwritten_02.png",
    "PL-015",
    "GO",
    "Sarah Johnson",
    "2025-01-15",
    None,
    20
)

create_handwritten_form_text(
    "form_handwritten_03.png",
    "PL-078",
    "NO-GO",
    "Michael Brown",
    "2025-01-15",
    "Maintenance required - scheduled downtime",
    20
)

# Digital forms
create_digital_form(
    "form_online_01.docx",
    "PL-001",
    "GO",
    "Emily Davis",
    "2025-01-15",
    None,
    20
)

create_digital_form(
    "form_online_02.docx",
    "PL-025",
    "NO-GO",
    "Robert Wilson",
    "2025-01-15",
    "Quality issue detected - production halted for inspection",
    20
)

create_digital_form(
    "form_online_03.docx",
    "PL-050",
    "GO",
    "Lisa Anderson",
    "2025-01-15",
    None,
    20
)

print(f"\nAll mock forms created in: {MOCK_DATA_DIR}")
print("You can now upload these files in the application!")

