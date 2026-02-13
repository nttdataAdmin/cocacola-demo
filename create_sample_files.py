"""
Script to create sample requirements documents for ATF testing
"""
from pathlib import Path
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("Warning: python-docx not available. Please install it: pip install python-docx")
    DOCX_AVAILABLE = False
    sys.exit(1)

# Create sample files directory
SAMPLES_DIR = Path(__file__).parent / "SampleInputs"
SAMPLES_DIR.mkdir(exist_ok=True)

def create_sample_1():
    """E-commerce Product Search Requirements"""
    doc = Document()
    
    # Title
    title = doc.add_heading('E-Commerce Product Search Feature', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    doc.add_paragraph('This document outlines the requirements for implementing an enhanced product search feature in the e-commerce platform.')
    
    # Section 1: Functional Requirements
    doc.add_heading('1. Functional Requirements', 1)
    
    doc.add_heading('1.1 Search Functionality', 2)
    doc.add_paragraph('The system shall provide a search interface that allows users to search for products using keywords.', style='List Bullet')
    doc.add_paragraph('The search shall support partial matching and fuzzy search capabilities.', style='List Bullet')
    doc.add_paragraph('The search results shall be displayed in a paginated format with 20 items per page.', style='List Bullet')
    doc.add_paragraph('Users shall be able to filter search results by category, brand, price range, and seller attributes.', style='List Bullet')
    
    doc.add_heading('1.2 Faceted Search', 2)
    doc.add_paragraph('The system shall display available facets (filters) based on the search results.', style='List Bullet')
    doc.add_paragraph('Seller Attribute shall be displayed as a facet in the Search Engine Product Listing Page (PLP).', style='List Bullet')
    doc.add_paragraph('Facets shall be dynamically updated based on the current search results.', style='List Bullet')
    
    doc.add_heading('1.3 Search Results Display', 2)
    doc.add_paragraph('Each search result shall display: Product Name, SKU ID, Category, Brand, Model Number, Description, Price, and Seller Information.', style='List Bullet')
    doc.add_paragraph('Results shall be sortable by: Relevance, Price (Low to High), Price (High to Low), and Newest First.', style='List Bullet')
    
    # Section 2: Non-Functional Requirements
    doc.add_heading('2. Non-Functional Requirements', 1)
    
    doc.add_paragraph('The search response time shall be less than 2 seconds for queries returning up to 10,000 results.', style='List Bullet')
    doc.add_paragraph('The system shall support concurrent searches from at least 1000 users.', style='List Bullet')
    doc.add_paragraph('The search functionality shall be accessible via web and mobile interfaces.', style='List Bullet')
    
    # Section 3: Business Objectives
    doc.add_heading('3. Business Objectives', 1)
    
    doc.add_paragraph('Improve user experience by providing fast and accurate search results.', style='List Bullet')
    doc.add_paragraph('Increase product discoverability through enhanced filtering options.', style='List Bullet')
    doc.add_paragraph('Enable sellers to showcase their products effectively through seller attribute filtering.', style='List Bullet')
    
    # Section 4: Dependencies
    doc.add_heading('4. Dependencies', 1)
    
    doc.add_paragraph('Product cATFlog dATFbase with indexed product information.', style='List Bullet')
    doc.add_paragraph('Search engine infrastructure (Elasticsearch or similar).', style='List Bullet')
    doc.add_paragraph('Seller management system for seller attribute dATF.', style='List Bullet')
    
    # Save
    output_file = SAMPLES_DIR / "ECommerce_ProductSearch_Requirements.docx"
    doc.save(output_file)
    print(f"Created: {output_file}")

def create_sample_2():
    """User Authentication and Authorization Requirements"""
    doc = Document()
    
    title = doc.add_heading('User Authentication and Authorization System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document specifies the requirements for implementing a secure user authentication and authorization system.')
    
    # Section 1: Authentication
    doc.add_heading('1. Authentication Requirements', 1)
    
    doc.add_heading('1.1 User Login', 2)
    doc.add_paragraph('Users shall be able to log in using email address and password.', style='List Bullet')
    doc.add_paragraph('The system shall support password reset functionality via email verification.', style='List Bullet')
    doc.add_paragraph('Failed login attempts shall be limited to 5 attempts within 15 minutes, after which the account shall be temporarily locked.', style='List Bullet')
    
    doc.add_heading('1.2 Multi-Factor Authentication', 2)
    doc.add_paragraph('The system shall support optional two-factor authentication (2FA) using SMS or authenticator apps.', style='List Bullet')
    doc.add_paragraph('Administrative users shall be required to use 2FA.', style='List Bullet')
    
    # Section 2: Authorization
    doc.add_heading('2. Authorization Requirements', 1)
    
    doc.add_paragraph('The system shall implement role-based access control (RBAC) with the following roles: Admin, Manager, User, Guest.', style='List Bullet')
    doc.add_paragraph('Users shall only access resources and perform actions permitted by their assigned role.', style='List Bullet')
    doc.add_paragraph('Permission changes shall take effect immediately without requiring user re-login.', style='List Bullet')
    
    # Section 3: Session Management
    doc.add_heading('3. Session Management', 1)
    
    doc.add_paragraph('User sessions shall expire after 30 minutes of inactivity.', style='List Bullet')
    doc.add_paragraph('Users shall be able to view and terminate active sessions from their account settings.', style='List Bullet')
    doc.add_paragraph('The system shall support "Remember Me" functionality with secure token-based authentication.', style='List Bullet')
    
    # Section 4: Security Requirements
    doc.add_heading('4. Security Requirements', 1)
    
    doc.add_paragraph('All passwords shall be hashed using bcrypt with salt.', style='List Bullet')
    doc.add_paragraph('All authentication endpoints shall use HTTPS only.', style='List Bullet')
    doc.add_paragraph('The system shall implement protection against SQL injection and XSS attacks.', style='List Bullet')
    
    # Save
    output_file = SAMPLES_DIR / "UserAuth_Authorization_Requirements.docx"
    doc.save(output_file)
    print(f"Created: {output_file}")

def create_sample_3():
    """Payment Processing System Requirements"""
    doc = Document()
    
    title = doc.add_heading('Payment Processing System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document describes the requirements for implementing a secure payment processing system for online transactions.')
    
    # Section 1: Payment Methods
    doc.add_heading('1. Payment Methods', 1)
    
    doc.add_paragraph('The system shall support the following payment methods: Credit Card, Debit Card, PayPal, Bank Transfer, and Digital Wallets.', style='List Bullet')
    doc.add_paragraph('Users shall be able to save multiple payment methods for future use.', style='List Bullet')
    doc.add_paragraph('Payment method information shall be encrypted and stored securely in compliance with PCI DSS standards.', style='List Bullet')
    
    # Section 2: Transaction Processing
    doc.add_heading('2. Transaction Processing', 1)
    
    doc.add_heading('2.1 Payment Authorization', 2)
    doc.add_paragraph('The system shall authorize payments before processing the transaction.', style='List Bullet')
    doc.add_paragraph('Authorization shall be verified with the payment gateway within 5 seconds.', style='List Bullet')
    doc.add_paragraph('Failed authorizations shall be logged and the user shall be notified immediately.', style='List Bullet')
    
    doc.add_heading('2.2 Payment Confirmation', 2)
    doc.add_paragraph('Upon successful payment, users shall receive an email confirmation with transaction details.', style='List Bullet')
    doc.add_paragraph('The system shall generate a unique transaction ID for each payment.', style='List Bullet')
    doc.add_paragraph('Payment status shall be updated in real-time in the user account.', style='List Bullet')
    
    # Section 3: Refund Processing
    doc.add_heading('3. Refund Processing', 1)
    
    doc.add_paragraph('Authorized users shall be able to initiate refunds for eligible transactions.', style='List Bullet')
    doc.add_paragraph('Refunds shall be processed within 5-7 business days.', style='List Bullet')
    doc.add_paragraph('The system shall support partial refunds for orders with multiple items.', style='List Bullet')
    
    # Section 4: Reporting
    doc.add_heading('4. Reporting and Analytics', 1)
    
    doc.add_paragraph('The system shall generate daily, weekly, and monthly payment reports.', style='List Bullet')
    doc.add_paragraph('Reports shall include: Total transactions, Success rate, Failed transactions, and Revenue by payment method.', style='List Bullet')
    
    # Save
    output_file = SAMPLES_DIR / "PaymentProcessing_Requirements.docx"
    doc.save(output_file)
    print(f"Created: {output_file}")

def create_sample_4():
    """Order Management System Requirements"""
    doc = Document()
    
    title = doc.add_heading('Order Management System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document outlines the requirements for an order management system to handle customer orders from placement to fulfillment.')
    
    # Section 1: Order Placement
    doc.add_heading('1. Order Placement', 1)
    
    doc.add_paragraph('Users shall be able to place orders by adding items to cart and proceeding to checkout.', style='List Bullet')
    doc.add_paragraph('The system shall validate product availability before order confirmation.', style='List Bullet')
    doc.add_paragraph('Orders shall be assigned a unique order number upon successful placement.', style='List Bullet')
    doc.add_paragraph('Users shall receive order confirmation email with order details and tracking information.', style='List Bullet')
    
    # Section 2: Order Status
    doc.add_heading('2. Order Status Management', 1)
    
    doc.add_paragraph('The system shall track order status through the following stages: Pending, Confirmed, Processing, Shipped, Delivered, Cancelled.', style='List Bullet')
    doc.add_paragraph('Order status updates shall be sent to users via email and SMS notifications.', style='List Bullet')
    doc.add_paragraph('Users shall be able to view order status and tracking information from their account dashboard.', style='List Bullet')
    
    # Section 3: Order Fulfillment
    doc.add_heading('3. Order Fulfillment', 1)
    
    doc.add_paragraph('The system shall generate shipping labels automatically upon order confirmation.', style='List Bullet')
    doc.add_paragraph('Inventory shall be reserved upon order confirmation and released if order is cancelled.', style='List Bullet')
    doc.add_paragraph('The system shall support multiple shipping methods: Standard, Express, and Overnight delivery.', style='List Bullet')
    
    # Section 4: Order Cancellation
    doc.add_heading('4. Order Cancellation', 1)
    
    doc.add_paragraph('Users shall be able to cancel orders within 24 hours of placement if order status is Pending or Confirmed.', style='List Bullet')
    doc.add_paragraph('Cancelled orders shall trigger automatic refund processing.', style='List Bullet')
    doc.add_paragraph('Administrators shall be able to cancel orders at any stage with appropriate reason logging.', style='List Bullet')
    
    # Save
    output_file = SAMPLES_DIR / "OrderManagement_Requirements.docx"
    doc.save(output_file)
    print(f"Created: {output_file}")

if __name__ == "__main__":
    print("Creating sample requirements documents...")
    print("=" * 50)
    
    create_sample_1()
    create_sample_2()
    create_sample_3()
    create_sample_4()
    
    print("=" * 50)
    print(f"\nAll sample files created in: {SAMPLES_DIR}")
    print("\nSample files created:")
    print("1. ECommerce_ProductSearch_Requirements.docx - E-commerce search feature")
    print("2. UserAuth_Authorization_Requirements.docx - Authentication system")
    print("3. PaymentProcessing_Requirements.docx - Payment processing")
    print("4. OrderManagement_Requirements.docx - Order management system")
    print("\nYou can now upload these files to test the ATF system!")

