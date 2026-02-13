from fastapi import FastAPI, HTTPException, UploadFile, File, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pathlib import Path
from pydantic import BaseModel
from typing import Optional, List
import json
import os
import zipfile
import io
from datetime import datetime
from openai import AzureOpenAI
from dotenv import load_dotenv
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. User story saving will be skipped.")
try:
    import pandas as pd
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("Warning: pandas/openpyxl not available. Excel export will be skipped.")
import asyncio

# Load environment variables from .env file
load_dotenv()

app = FastAPI()

# Azure OpenAI Configuration - Load from environment variables
AZURE_ENDPOINT = os.getenv("AZURE_ENDPOINT", "https://oai-mcm-agentic-flow-nprd01.openai.azure.com/")
AZURE_DEPLOYMENT = os.getenv("AZURE_DEPLOYMENT", "gpt-4.1")
AZURE_API_KEY = os.getenv("AZURE_API_KEY", "")
AZURE_API_VERSION = os.getenv("AZURE_API_VERSION", "2024-08-01-preview")

# Output directories
BASE_DIR = Path(__file__).parent
OUTPUT_DIR = BASE_DIR / "Output"
USER_STORY_DIR = OUTPUT_DIR / "UserStory"
TEST_CASES_DIR = OUTPUT_DIR / "TestCases"
TEST_DATA_DIR = OUTPUT_DIR / "TestData"
TEST_COVERAGE_DIR = OUTPUT_DIR / "TestCoverage"
EXCEL_OUTPUT_DIR = OUTPUT_DIR / "Excel"

# Create output directories if they don't exist
for dir_path in [OUTPUT_DIR, USER_STORY_DIR, TEST_CASES_DIR, TEST_DATA_DIR, TEST_COVERAGE_DIR, EXCEL_OUTPUT_DIR]:
    dir_path.mkdir(parents=True, exist_ok=True)

# Enable CORS for React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Store uploaded files in memory
uploaded_files = {}

def get_azure_client():
    """Get Azure OpenAI client"""
    return AzureOpenAI(
        azure_endpoint=AZURE_ENDPOINT,
        api_key=AZURE_API_KEY,
        api_version=AZURE_API_VERSION
    )

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        if not DOCX_AVAILABLE:
            return "DOCX parsing not available. Please install python-docx."
        from io import BytesIO
        doc = Document(BytesIO(file_content))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return ""

@app.post("/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    """Upload DOCX files"""
    try:
        uploaded_data = []
        for file in files:
            if not file.filename.endswith('.docx'):
                continue
            
            content = await file.read()
            text = extract_text_from_docx(content)
            
            file_data = {
                "filename": file.filename,
                "content": text,
                "size": len(content),
                "uploaded_at": datetime.now().isoformat()
            }
            uploaded_data.append(file_data)
            uploaded_files[file.filename] = file_data
        
        return {
            "success": True,
            "files": uploaded_data,
            "count": len(uploaded_data),
            "message": f"Successfully uploaded {len(uploaded_data)} file(s)"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/uploaded-files")
async def get_uploaded_files():
    """Get list of uploaded files"""
    return {
        "files": list(uploaded_files.values()),
        "count": len(uploaded_files)
    }

@app.post("/agent/requirements-analyst")
async def requirements_analyst_agent(file_name: str):
    """Agent 1: Requirements Analyst - Analyze requirements document"""
    try:
        if file_name not in uploaded_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        file_data = uploaded_files[file_name]
        requirements_text = file_data["content"]
        
        # Validate that we have content
        if not requirements_text or len(requirements_text.strip()) == 0:
            raise HTTPException(status_code=400, detail="Document content is empty. Please ensure the DOCX file contains text.")
        
        print(f"Analyzing document: {file_name}, Content length: {len(requirements_text)} characters")
        
        client = get_azure_client()
        
        prompt = f"""You are a Requirements Analyst. You have been given a requirements document below. Your task is to ANALYZE THIS DOCUMENT and provide a comprehensive overview.

=== REQUIREMENTS DOCUMENT TO ANALYZE ===
{requirements_text}
=== END OF DOCUMENT ===

CRITICAL INSTRUCTIONS:
- The document above is the ACTUAL requirements document you need to analyze
- DO NOT ask for the document - it is already provided above
- You MUST analyze the content provided and extract information from it
- If the document appears empty or minimal, analyze what is there

TASK - Analyze the document above and provide:

1. SUMMARY (2-3 paragraphs): Write a comprehensive overview that explains:
   - What this document is about
   - What the system/application needs to do based on the document
   - The main purpose and goals mentioned in the document
   - Key business needs and objectives from the document
   - Overall scope and context described in the document

2. EXTRACT the following from the document:
   - All functional requirements (what the system must do)
   - Key features and functionalities mentioned
   - Business objectives stated
   - Non-functional requirements (performance, security, etc.)
   - Dependencies mentioned
   - Constraints or limitations described

OUTPUT FORMAT - Provide your analysis in JSON format ONLY:
{{
    "summary": "A comprehensive 2-3 paragraph overview analyzing the document above. Explain what the document contains, what the system needs to accomplish based on the document content, the main business objectives from the document, and the overall scope. Write in clear, professional language.",
    "functional_requirements": ["list of functional requirements extracted from the document"],
    "key_features": ["list of key features extracted from the document"],
    "business_objectives": ["list of business objectives from the document"],
    "non_functional_requirements": ["list of non-functional requirements from the document"],
    "dependencies": ["list of dependencies mentioned in the document"],
    "constraints": ["list of constraints from the document"]
}}

REMEMBER: Analyze the document provided above. Do not ask for the document - it is already given to you."""

        response = client.chat.completions.create(
            model=AZURE_DEPLOYMENT,
            messages=[
                {"role": "system", "content": "You are an expert Requirements Analyst with deep expertise in software requirements analysis, business analysis, and system design. You excel at understanding complex requirements documents and providing clear, comprehensive overviews that help stakeholders understand what needs to be built and why."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=4000
        )
        
        result_text = response.choices[0].message.content
        
        # Clean and parse JSON - handle markdown code blocks
        try:
            # Remove markdown code blocks if present
            cleaned_text = result_text.strip()
            if "```json" in cleaned_text:
                cleaned_text = cleaned_text.split("```json")[1].split("```")[0].strip()
            elif "```" in cleaned_text:
                cleaned_text = cleaned_text.split("```")[1].split("```")[0].strip()
            
            result_json = json.loads(cleaned_text)
        except json.JSONDecodeError as e:
            print(f"JSON parsing error in requirements analyst: {e}")
            print(f"Response preview: {result_text[:500]}")
            # If not JSON, create structured response
            result_json = {
                "analysis": result_text,
                "raw_response": result_text
            }
        except Exception as e:
            print(f"Error parsing requirements analysis: {e}")
            result_json = {
                "analysis": result_text,
                "raw_response": result_text
            }
        
        result_json["agent_name"] = "Requirements Analyst"
        result_json["file_name"] = file_name
        result_json["timestamp"] = datetime.now().isoformat()
        
        return result_json
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/agent/user-story-creator")
async def user_story_creator_agent(file_name: str):
    """Agent 2: User Story Creator - Generate user stories from requirements"""
    try:
        if file_name not in uploaded_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        file_data = uploaded_files[file_name]
        requirements_text = file_data["content"]
        
        client = get_azure_client()
        
        prompt = f"""You are a User Story Creator. Based on the following requirements document, create comprehensive user stories following the standard format: "As a [user type], I want [functionality] so that [benefit]."

REQUIREMENTS DOCUMENT:
{requirements_text}

TASK:
1. Create AT LEAST 15-20 user stories for all functional requirements (more if the document is complex)
2. Break down each major feature into multiple user stories
3. Include acceptance criteria for each user story (minimum 3-5 criteria per story)
4. Format each user story with:
   - Title
   - User Story (As a... I want... So that...)
   - Acceptance Criteria (list of criteria)
   - Priority (High/Medium/Low)

IMPORTANT: Generate comprehensive user stories covering ALL features mentioned in the requirements. Do not skip any functionality.

Provide the output in JSON format ONLY (no markdown, no explanations, just valid JSON):
{{
    "user_stories": [
        {{
            "id": "US-001",
            "title": "Story title",
            "user_story": "As a [user type], I want [functionality] so that [benefit]",
            "acceptance_criteria": ["criterion 1", "criterion 2", "criterion 3"],
            "priority": "High"
        }}
    ],
    "summary": "Summary of generated user stories"
}}"""

        response = client.chat.completions.create(
            model=AZURE_DEPLOYMENT,
            messages=[
                {"role": "system", "content": "You are an expert in creating user stories from requirements."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=8000
        )
        
        result_text = response.choices[0].message.content
        
        # Clean and parse JSON - handle markdown code blocks
        try:
            # Remove markdown code blocks if present
            cleaned_text = result_text.strip()
            if "```json" in cleaned_text:
                cleaned_text = cleaned_text.split("```json")[1].split("```")[0].strip()
            elif "```" in cleaned_text:
                cleaned_text = cleaned_text.split("```")[1].split("```")[0].strip()
            
            result_json = json.loads(cleaned_text)
            
            # Validate that we have user stories
            if not result_json.get("user_stories") or len(result_json.get("user_stories", [])) == 0:
                print(f"Warning: No user stories found in parsed JSON. Raw response length: {len(result_text)}")
                result_json["raw_response"] = result_text
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {e}")
            print(f"Response preview: {result_text[:500]}")
            result_json = {
                "user_stories": [],
                "raw_response": result_text
            }
        except Exception as e:
            print(f"Error parsing user stories: {e}")
            result_json = {
                "user_stories": [],
                "raw_response": result_text
            }
        
        result_json["agent_name"] = "User Story Creator"
        result_json["file_name"] = file_name
        result_json["timestamp"] = datetime.now().isoformat()
        
        # Parse user stories from raw_response if array is empty
        user_stories_list = result_json.get("user_stories", [])
        if not user_stories_list and "raw_response" in result_json:
            try:
                raw_text = result_json["raw_response"]
                if "```json" in raw_text:
                    raw_text = raw_text.split("```json")[1].split("```")[0].strip()
                elif "```" in raw_text:
                    raw_text = raw_text.split("```")[1].split("```")[0].strip()
                parsed = json.loads(raw_text)
                user_stories_list = parsed.get("user_stories", [])
                result_json["user_stories"] = user_stories_list  # Update the result
                print(f"Parsed {len(user_stories_list)} user stories from raw_response")
            except Exception as e:
                print(f"Error parsing user stories from raw_response: {e}")
        
        # Save user story to file
        output_file = USER_STORY_DIR / f"{Path(file_name).stem}_userstory.docx"
        try:
            if DOCX_AVAILABLE and user_stories_list:
                doc = Document()
                doc.add_heading('User Stories', 0)
                for story in user_stories_list:
                    doc.add_heading(story.get("title", "Untitled"), level=1)
                    doc.add_paragraph(f"ID: {story.get('id', 'N/A')}")
                    doc.add_paragraph(f"User Story: {story.get('user_story', 'N/A')}")
                    doc.add_paragraph("Acceptance Criteria:")
                    for criteria in story.get("acceptance_criteria", []):
                        doc.add_paragraph(criteria, style='List Bullet')
                    doc.add_paragraph(f"Priority: {story.get('priority', 'N/A')}")
                    doc.add_paragraph("")
                doc.save(output_file)
                print(f"Saved {len(user_stories_list)} user stories to DOCX")
            # Always save as JSON (backup and for empty cases)
            json_file = USER_STORY_DIR / f"{Path(file_name).stem}_userstory.json"
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(result_json, f, indent=2)
        except Exception as e:
            print(f"Error saving user story: {e}")
            # Fallback: save as JSON
            try:
                json_file = USER_STORY_DIR / f"{Path(file_name).stem}_userstory.json"
                with open(json_file, 'w', encoding='utf-8') as f:
                    json.dump(result_json, f, indent=2)
            except Exception as e2:
                print(f"Error saving user story as JSON: {e2}")
        
        return result_json
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/agent/test-case-generator")
async def test_case_generator_agent(file_name: str, user_stories: dict = None):
    """Agent 3: Test Case Generator - Generate test cases from user stories"""
    try:
        if file_name not in uploaded_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        file_data = uploaded_files[file_name]
        requirements_text = file_data["content"]
        
        client = get_azure_client()
        
        # Include user stories if provided
        user_stories_text = ""
        if user_stories:
            user_stories_text = f"\n\nUSER STORIES:\n{json.dumps(user_stories, indent=2)}"
        
        prompt = f"""You are a Test Case Generator. Based on the requirements and user stories, create comprehensive test cases.

REQUIREMENTS DOCUMENT:
{requirements_text}
{user_stories_text}

TASK:
1. Create AT LEAST 50-100 test cases covering all user stories and requirements
2. For each user story, generate 3-5 test cases (positive, negative, boundary, edge cases)
3. Each test case should include:
   - Test Case ID (sequential: TC-001, TC-002, etc.)
   - Test Case Title
   - Description
   - Preconditions (list)
   - Test Steps (detailed step-by-step list)
   - Expected Results
   - Priority (High/Medium/Low)
   - Test Type (Functional/Non-functional)
   - User Story ID (link to user story)

IMPORTANT: 
- Generate comprehensive test coverage (minimum 50 test cases)
- Include positive, negative, boundary, and edge case scenarios
- Cover all functionalities mentioned in requirements
- Be thorough and detailed

Provide the output in JSON format ONLY (no markdown, no explanations, just valid JSON):
{{
    "test_cases": [
        {{
            "id": "TC-001",
            "title": "Test case title",
            "description": "Test case description",
            "preconditions": ["precondition 1", "precondition 2"],
            "test_steps": ["step 1", "step 2", "step 3"],
            "expected_results": "Expected result description",
            "priority": "High",
            "test_type": "Functional",
            "user_story_id": "US-001"
        }}
    ],
    "summary": "Summary of generated test cases"
}}"""

        response = client.chat.completions.create(
            model=AZURE_DEPLOYMENT,
            messages=[
                {"role": "system", "content": "You are an expert in creating comprehensive test cases."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=16000
        )
        
        result_text = response.choices[0].message.content
        
        # Clean and parse JSON - handle markdown code blocks
        try:
            # Remove markdown code blocks if present
            cleaned_text = result_text.strip()
            if "```json" in cleaned_text:
                cleaned_text = cleaned_text.split("```json")[1].split("```")[0].strip()
            elif "```" in cleaned_text:
                cleaned_text = cleaned_text.split("```")[1].split("```")[0].strip()
            
            result_json = json.loads(cleaned_text)
            
            # Validate that we have test cases
            if not result_json.get("test_cases") or len(result_json.get("test_cases", [])) == 0:
                print(f"Warning: No test cases found in parsed JSON. Raw response length: {len(result_text)}")
                result_json["raw_response"] = result_text
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {e}")
            print(f"Response preview: {result_text[:500]}")
            result_json = {
                "test_cases": [],
                "raw_response": result_text
            }
        except Exception as e:
            print(f"Error parsing test cases: {e}")
            result_json = {
                "test_cases": [],
                "raw_response": result_text
            }
        
        result_json["agent_name"] = "Test Case Generator"
        result_json["file_name"] = file_name
        result_json["timestamp"] = datetime.now().isoformat()
        
        return result_json
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/agent/test-data-generator")
async def test_data_generator_agent(file_name: str, test_cases: dict = None):
    """Agent 4: Test Data Generator - Generate Test Data for test cases"""
    try:
        if file_name not in uploaded_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        file_data = uploaded_files[file_name]
        requirements_text = file_data["content"]
        
        client = get_azure_client()
        
        test_cases_text = ""
        if test_cases:
            test_cases_text = f"\n\nTEST CASES:\n{json.dumps(test_cases, indent=2)}"
        
        prompt = f"""You are a Test Data Generator. Generate comprehensive Test Data in JSON format for the test cases.

REQUIREMENTS DOCUMENT:
{requirements_text}
{test_cases_text}

TASK:
1. Generate Test Data for EACH test case (minimum 2-3 data sets per test case: positive, negative, boundary)
2. Include realistic, diverse data values
3. For e-commerce/product scenarios, include: SKU_ID, Product_name, Category, Brand, Model_Number, Description, Price, Seller, etc.
4. For other scenarios, include all relevant data fields based on the requirements
5. Generate comprehensive Test Data covering all scenarios
6. Ensure data sets are realistic and varied

CRITICAL FORMAT REQUIREMENT:
- Use a FLAT structure (like test cases format) - DO NOT nest data in a "data" object
- All data fields should be at the top level of each Test Data object
- This makes it easier to export to Excel and fetch individual fields

IMPORTANT:
- Generate Test Data for ALL test cases provided
- Create multiple data variations (positive, negative, boundary, edge cases)
- Make data realistic and comprehensive

Provide the output in JSON format ONLY (no markdown, no explanations, just valid JSON):
{{
    "test_data": [
        {{
            "id": "TD-001",
            "test_case_id": "TC-001",
            "data_set_name": "Positive Test Data",
            "SKU_ID": "SKU123",
            "Product_name": "Sample Product",
            "Category": "Electronics",
            "Brand": "Sample Brand",
            "Model_Number": "MOD-001",
            "Description": "Product description",
            "Price": 99.99,
            "description": "Test Data description"
        }},
        {{
            "id": "TD-002",
            "test_case_id": "TC-001",
            "data_set_name": "Negative Test Data",
            "SKU_ID": "",
            "Product_name": "",
            "Category": "Electronics",
            "Brand": "Sample Brand",
            "Model_Number": "MOD-001",
            "Description": "Missing required fields",
            "Price": 0,
            "description": "Test Data with missing mandatory fields"
        }}
    ],
    "summary": "Summary of generated Test Data"
}}

REMEMBER: Use FLAT structure - all data fields at top level, NOT nested in a "data" object."""

        response = client.chat.completions.create(
            model=AZURE_DEPLOYMENT,
            messages=[
                {"role": "system", "content": "You are an expert in generating Test Data for software testing."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=12000
        )
        
        result_text = response.choices[0].message.content
        
        # Clean and parse JSON - handle markdown code blocks
        try:
            # Remove markdown code blocks if present
            cleaned_text = result_text.strip()
            if "```json" in cleaned_text:
                cleaned_text = cleaned_text.split("```json")[1].split("```")[0].strip()
            elif "```" in cleaned_text:
                cleaned_text = cleaned_text.split("```")[1].split("```")[0].strip()
            
            result_json = json.loads(cleaned_text)
            
            # Validate that we have Test Data
            if not result_json.get("test_data") or len(result_json.get("test_data", [])) == 0:
                print(f"Warning: No Test Data found in parsed JSON. Raw response length: {len(result_text)}")
                result_json["raw_response"] = result_text
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {e}")
            print(f"Response preview: {result_text[:500]}")
            result_json = {
                "test_data": [],
                "raw_response": result_text
            }
        except Exception as e:
            print(f"Error parsing Test Data: {e}")
            result_json = {
                "test_data": [],
                "raw_response": result_text
            }
        
        result_json["agent_name"] = "Test Data Generator"
        result_json["file_name"] = file_name
        result_json["timestamp"] = datetime.now().isoformat()
        
        # Save Test Data to JSON file
        output_file = TEST_DATA_DIR / f"{Path(file_name).stem}_testdata.json"
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(result_json, f, indent=2)
        except Exception as e:
            print(f"Error saving Test Data: {e}")
        
        return result_json
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def export_to_excel(file_name: str, test_cases: dict, test_data: dict) -> Path:
    """Export test cases and Test Data to Excel file"""
    if not EXCEL_AVAILABLE:
        raise Exception("Excel libraries not available")
    
    file_stem = Path(file_name).stem
    excel_file = EXCEL_OUTPUT_DIR / f"{file_stem}_TestCases_TestData.xlsx"
    
    wb = Workbook()
    
    # Test Cases Sheet
    ws_testcases = wb.active
    ws_testcases.title = "Test Cases"
    
    # Headers for Test Cases
    headers_tc = ["Test Case ID", "Test Case Name", "Description", "Preconditions", "Test Steps", "Expected Result", "Priority", "Status"]
    ws_testcases.append(headers_tc)
    
    # Style headers
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    for cell in ws_testcases[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Parse test cases - check if they're in raw_response
    test_cases_list = test_cases.get("test_cases", [])
    
    # If test_cases is empty, try to parse raw_response
    if not test_cases_list and "raw_response" in test_cases:
        try:
            raw_response = test_cases["raw_response"]
            # Remove markdown code blocks if present
            if "```json" in raw_response:
                raw_response = raw_response.split("```json")[1].split("```")[0].strip()
            elif "```" in raw_response:
                raw_response = raw_response.split("```")[1].split("```")[0].strip()
            
            parsed_data = json.loads(raw_response)
            test_cases_list = parsed_data.get("test_cases", [])
        except (json.JSONDecodeError, KeyError) as e:
            print(f"Error parsing raw_response: {e}")
            test_cases_list = []
    
    # Add test cases data
    for idx, tc in enumerate(test_cases_list, start=2):
        # Handle different field name variations
        test_case_id = tc.get("test_case_id") or tc.get("id", f"TC_{idx-1}")
        test_case_name = tc.get("test_case_name") or tc.get("title", "")
        description = tc.get("description", "")
        
        # Handle preconditions - can be string, list, or dict
        preconditions = tc.get("preconditions", "")
        if isinstance(preconditions, list):
            preconditions = "\n".join([str(p) for p in preconditions])
        elif isinstance(preconditions, dict):
            preconditions = "\n".join([f"{k}: {v}" for k, v in preconditions.items()])
        else:
            preconditions = str(preconditions)
        
        # Handle test steps
        test_steps = tc.get("test_steps", [])
        if isinstance(test_steps, list):
            test_steps_str = "\n".join([f"{i+1}. {step}" if isinstance(step, str) else str(step) for i, step in enumerate(test_steps)])
        else:
            test_steps_str = str(test_steps)
        
        # Handle expected result - can be expected_result or expected_results
        expected_result = tc.get("expected_result") or tc.get("expected_results", "")
        
        priority = tc.get("priority", "Medium")
        status = tc.get("status", "Not Executed")
        
        ws_testcases.append([
            test_case_id,
            test_case_name,
            description,
            preconditions,
            test_steps_str,
            expected_result,
            priority,
            status
        ])
    
    # Auto-adjust column widths
    for column in ws_testcases.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws_testcases.column_dimensions[column_letter].width = adjusted_width
    
    # Test Data Sheet
    ws_testdata = wb.create_sheet("Test Data")
    
    # Parse Test Data - check if they're in raw_response (must parse first to collect fields)
    test_data_list = test_data.get("test_data", [])
    
    # If test_data is empty, try to parse raw_response
    if not test_data_list and "raw_response" in test_data:
        try:
            raw_response = test_data["raw_response"]
            
            # Handle case where raw_response is a JSON string (double-encoded)
            if isinstance(raw_response, str) and raw_response.strip().startswith('{'):
                try:
                    # First, try to parse it as a JSON string
                    parsed_first = json.loads(raw_response)
                    # If the result is still a string, parse again
                    if isinstance(parsed_first, str):
                        raw_response = parsed_first
                    elif isinstance(parsed_first, dict):
                        # If it's already a dict, check if it has test_data
                        test_data_list = parsed_first.get("test_data", [])
                        if test_data_list:
                            print(f"Found {len(test_data_list)} Test Data sets from parsed raw_response")
                except json.JSONDecodeError:
                    pass  # Continue with original raw_response
            
            # Remove markdown code blocks if present
            if isinstance(raw_response, str):
                if "```json" in raw_response:
                    raw_response = raw_response.split("```json")[1].split("```")[0].strip()
                elif "```" in raw_response:
                    raw_response = raw_response.split("```")[1].split("```")[0].strip()
                
                # Try to parse the JSON
                if not test_data_list:
                    parsed_data = json.loads(raw_response)
                    test_data_list = parsed_data.get("test_data", [])
                    if test_data_list:
                        print(f"Found {len(test_data_list)} Test Data sets from raw_response after markdown removal")
        except (json.JSONDecodeError, KeyError) as e:
            print(f"Error parsing Test Data raw_response: {e}")
            print(f"Raw response type: {type(test_data.get('raw_response'))}")
            print(f"Raw response preview: {str(test_data.get('raw_response', ''))[:200]}")
            test_data_list = []
    
    print(f"Total Test Data sets to export: {len(test_data_list)}")
    
    # Collect all unique data field names from Test Data (for flat structure)
    # This must be done AFTER parsing test_data_list
    all_data_fields = set()
    for td in test_data_list:
        if isinstance(td, dict):
            # Get all fields except standard ones (new flat format)
            standard_fields = {"id", "test_data_id", "test_case_id", "data_set_name", "name", "description", "data", "data_values"}
            for key in td.keys():
                if key not in standard_fields:
                    all_data_fields.add(key)
            # Also check nested data field (old format for backward compatibility)
            if "data" in td and isinstance(td["data"], dict):
                all_data_fields.update(td["data"].keys())
            if "data_values" in td and isinstance(td["data_values"], dict):
                all_data_fields.update(td["data_values"].keys())
    
    # Headers for Test Data - using flat structure similar to test cases
    headers_td = ["Test Data ID", "Test Case ID", "Data Set Name", "Description"]
    
    # Add data field columns (sorted for consistency)
    if all_data_fields:
        headers_td.extend(sorted(all_data_fields))
    else:
        # Fallback: if no fields found, use generic "Test Data Values"
        headers_td.append("Test Data Values")
    
    ws_testdata.append(headers_td)
    
    # Style headers
    for cell in ws_testdata[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
    
    # Add Test Data rows
    for idx, td in enumerate(test_data_list, start=2):
        # Handle different field name variations
        test_data_id = td.get("test_data_id") or td.get("id", f"TD_{idx-1}")
        test_case_id = td.get("test_case_id", "")
        data_set_name = td.get("data_set_name") or td.get("name", "")
        description = td.get("description", "")
        
        # Start row with standard fields
        row_data = [test_data_id, test_case_id, data_set_name, description]
        
        # Handle both flat structure (new) and nested structure (old format for backward compatibility)
        # Check if it's new flat format or old nested format
        standard_fields = {"id", "test_data_id", "test_case_id", "data_set_name", "name", "description", "data", "data_values"}
        data_fields_in_td = {k: v for k, v in td.items() if k not in standard_fields}
        
        if data_fields_in_td:
            # New format: data fields are flat at top level
            # Add each data field value in the same order as headers
            for field_name in sorted(all_data_fields) if all_data_fields else []:
                if field_name in data_fields_in_td:
                    value = data_fields_in_td[field_name]
                    # Format complex types
                    if isinstance(value, (dict, list)):
                        row_data.append(json.dumps(value))
                    else:
                        row_data.append(str(value))
                else:
                    row_data.append("")  # Empty cell if field not present in this row
        else:
            # Old format: data is nested in "data" or "data_values" field
            data_values = td.get("data_values") or td.get("data", {})
            
            if isinstance(data_values, dict):
                # Add each data field value in the same order as headers
                for field_name in sorted(all_data_fields) if all_data_fields else []:
                    if field_name in data_values:
                        value = data_values[field_name]
                        if isinstance(value, (dict, list)):
                            row_data.append(json.dumps(value))
                        else:
                            row_data.append(str(value))
                    else:
                        row_data.append("")  # Empty cell if field not present
            elif isinstance(data_values, list):
                # List format - convert to string
                row_data.append("\n".join([str(item) for item in data_values]))
            else:
                # Fallback: single value or string
                row_data.append(str(data_values) if data_values else "")
        
        ws_testdata.append(row_data)
    
    # Auto-adjust column widths for Test Data
    for column in ws_testdata.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws_testdata.column_dimensions[column_letter].width = adjusted_width
    
    wb.save(excel_file)
    return excel_file

class RunAgentsRequest(BaseModel):
    file_name: str

@app.post("/run-all-agents")
async def run_all_agents(request: RunAgentsRequest):
    """Run all agents in sequence for a file"""
    try:
        file_name = request.file_name
        if file_name not in uploaded_files:
            raise HTTPException(status_code=404, detail="File not found")
        
        results = {}
        
        # Agent 1: Requirements Analyst
        req_result = await requirements_analyst_agent(file_name)
        results["1_requirements_analyst"] = req_result
        
        # Agent 2: User Story Creator
        user_story_result = await user_story_creator_agent(file_name)
        results["2_user_story_creator"] = user_story_result
        
        # Agent 3: Test Case Generator
        test_case_result = await test_case_generator_agent(file_name, user_story_result)
        results["3_test_case_generator"] = test_case_result
        
        # Agent 4: Test Data Generator
        test_data_result = await test_data_generator_agent(file_name, test_case_result)
        results["4_test_data_generator"] = test_data_result
        
        # Parse user stories if in raw_response
        user_stories_list = user_story_result.get("user_stories", [])
        if not user_stories_list and "raw_response" in user_story_result:
            try:
                raw_text = user_story_result["raw_response"]
                if "```json" in raw_text:
                    raw_text = raw_text.split("```json")[1].split("```")[0].strip()
                elif "```" in raw_text:
                    raw_text = raw_text.split("```")[1].split("```")[0].strip()
                parsed = json.loads(raw_text)
                user_stories_list = parsed.get("user_stories", [])
            except:
                pass
        
        # Parse test cases if in raw_response
        test_cases_list = test_case_result.get("test_cases", [])
        if not test_cases_list and "raw_response" in test_case_result:
            try:
                raw_text = test_case_result["raw_response"]
                if "```json" in raw_text:
                    raw_text = raw_text.split("```json")[1].split("```")[0].strip()
                elif "```" in raw_text:
                    raw_text = raw_text.split("```")[1].split("```")[0].strip()
                parsed = json.loads(raw_text)
                test_cases_list = parsed.get("test_cases", [])
            except:
                pass
        
        # Parse Test Data if in raw_response - handle double-encoded JSON
        test_data_list = test_data_result.get("test_data", [])
        if not test_data_list and "raw_response" in test_data_result:
            try:
                raw_text = test_data_result["raw_response"]
                
                # Handle case where raw_response is a JSON string (double-encoded)
                if isinstance(raw_text, str) and raw_text.strip().startswith('{'):
                    try:
                        # First, try to parse it as a JSON string
                        parsed_first = json.loads(raw_text)
                        # If the result is still a string, parse again
                        if isinstance(parsed_first, str):
                            raw_text = parsed_first
                        elif isinstance(parsed_first, dict):
                            # If it's already a dict, check if it has test_data
                            test_data_list = parsed_first.get("test_data", [])
                            if test_data_list:
                                print(f"Found {len(test_data_list)} Test Data sets from parsed raw_response (first parse)")
                                # Update the result so it's available for Excel export
                                test_data_result["test_data"] = test_data_list
                    except json.JSONDecodeError:
                        pass  # Continue with original raw_response
                
                # If still empty, try removing markdown code blocks
                if not test_data_list and isinstance(raw_text, str):
                    if "```json" in raw_text:
                        raw_text = raw_text.split("```json")[1].split("```")[0].strip()
                    elif "```" in raw_text:
                        raw_text = raw_text.split("```")[1].split("```")[0].strip()
                    
                    # Try to parse the JSON
                    try:
                        parsed = json.loads(raw_text)
                        test_data_list = parsed.get("test_data", [])
                        if test_data_list:
                            print(f"Found {len(test_data_list)} Test Data sets from raw_response (after markdown removal)")
                            # Update the result so it's available for Excel export
                            test_data_result["test_data"] = test_data_list
                    except json.JSONDecodeError as e:
                        print(f"Error parsing Test Data from raw_response: {e}")
            except Exception as e:
                print(f"Error parsing Test Data: {e}")
        
        # Generate test coverage summary
        coverage = {
            "file_name": file_name,
            "timestamp": datetime.now().isoformat(),
            "user_stories_count": len(user_stories_list),
            "test_cases_count": len(test_cases_list),
            "test_data_sets_count": len(test_data_list),
            "coverage_percentage": 95  # Mock coverage
        }
        results["test_coverage"] = coverage
        
        # Save test coverage
        coverage_file = TEST_COVERAGE_DIR / f"{Path(file_name).stem}_testcoverage.json"
        try:
            with open(coverage_file, 'w', encoding='utf-8') as f:
                json.dump(coverage, f, indent=2)
        except Exception as e:
            print(f"Error saving test coverage: {e}")
        
        # Save test cases to file
        test_case_file = TEST_CASES_DIR / f"{Path(file_name).stem}_testcases.json"
        try:
            with open(test_case_file, 'w', encoding='utf-8') as f:
                json.dump(test_case_result, f, indent=2)
        except Exception as e:
            print(f"Error saving test cases: {e}")
        
        # Export to Excel
        if EXCEL_AVAILABLE:
            try:
                excel_file = export_to_excel(file_name, test_case_result, test_data_result)
                results["excel_file"] = str(excel_file)
            except Exception as e:
                print(f"Error exporting to Excel: {e}")
        
        return {
            "success": True,
            "file_name": file_name,
            "results": results,
            "timestamp": datetime.now().isoformat()
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class ChatRequest(BaseModel):
    message: str
    file_context: dict = None
    agent_results: dict = None
    chat_history: list = []

def get_ATF_knowledge_base(file_name: str, agent_results: dict) -> str:
    """Build comprehensive knowledge base from agent results"""
    kb = []
    
    if file_name in uploaded_files:
        kb.append(f"=== UPLOADED FILE ===\nFile: {file_name}\n")
    
    # Requirements Analyst
    if agent_results.get('1_requirements_analyst'):
        req = agent_results['1_requirements_analyst']
        kb.append(f"=== REQUIREMENTS ANALYSIS ===\n{json.dumps(req, indent=2)}\n")
    
    # User Story Creator
    if agent_results.get('2_user_story_creator'):
        us = agent_results['2_user_story_creator']
        kb.append(f"=== USER STORIES ===\n{json.dumps(us, indent=2)}\n")
    
    # Test Case Generator
    if agent_results.get('3_test_case_generator'):
        tc = agent_results['3_test_case_generator']
        kb.append(f"=== TEST CASES ===\n{json.dumps(tc, indent=2)}\n")
    
    # Test Data Generator
    if agent_results.get('4_test_data_generator'):
        td = agent_results['4_test_data_generator']
        kb.append(f"=== Test Data ===\n{json.dumps(td, indent=2)}\n")
    
    # Selenium Engineer
    if agent_results.get('5_selenium_engineer'):
        se = agent_results['5_selenium_engineer']
        kb.append(f"=== SELENIUM SCRIPTS ===\nScript generated: {se.get('script', 'N/A')[:500]}...\n")
    
    return "\n".join(kb)

@app.post("/chat")
async def chat_endpoint(request: ChatRequest):
    """AI Chat endpoint for ATF-related questions"""
    try:
        client = get_azure_client()
        
        # Build knowledge base
        file_name = request.file_context.get('file_name', 'Unknown') if request.file_context else 'No file selected'
        agent_results = request.agent_results or {}
        
        knowledge_base = get_ATF_knowledge_base(file_name, agent_results)
        
        system_message = f"""You are an Autonomous Testing Framework (ATF) assistant. You help users understand the testing workflow, requirements, user stories, test cases, Test Data, and Selenium scripts.

CURRENT FILE: {file_name}

KNOWLEDGE BASE:
{knowledge_base}

INSTRUCTIONS:
1. Answer questions based ONLY on the provided knowledge base
2. Be concise and professional (2-4 sentences when possible)
3. If asked about specific values, quote exact values from the knowledge base
4. If information is not available, say "This information is not available in the current context"
5. For technical questions about test cases or scripts, provide specific details"""
        
        messages = [{"role": "system", "content": system_message}]
        
        # Add chat history
        for msg in request.chat_history[-5:]:
            messages.append({"role": msg.get("role", "user"), "content": msg.get("content", "")})
        
        messages.append({"role": "user", "content": request.message})
        
        response = client.chat.completions.create(
            model=AZURE_DEPLOYMENT,
            messages=messages,
            temperature=0.3,
            max_tokens=500
        )
        
        return {"response": response.choices[0].message.content}
    except Exception as e:
        print(f"Chat error: {e}")
        return {"response": f"Error: {str(e)}"}

@app.get("/download-artifact/automation-design/{file_name:path}")
async def download_automation_design(file_name: str):
    """Download automation design artifact (.tsu file)"""
    import urllib.parse
    file_name = urllib.parse.unquote(file_name)
    
    # Use the P2P E2E Test.tsu file from Test Design Agent Output directory
    artifact_file = BASE_DIR / "Test Design Agent Output" / "P2P E2E Test.tsu"
    
    print(f"Looking for artifact file at: {artifact_file}")
    print(f"File exists: {artifact_file.exists()}")
    print(f"Absolute path: {artifact_file.absolute()}")
    
    if not artifact_file.exists():
        print(f"Error: File not found at {artifact_file.absolute()}")
        raise HTTPException(status_code=404, detail=f"Automation design artifact not found at {artifact_file.absolute()}")
    
    return FileResponse(
        path=str(artifact_file),
        filename=f"{file_name.replace('.docx', '')}_Automation_Design.tsu",
        media_type="application/octet-stream",
        headers={
            "Content-Disposition": f"attachment; filename={file_name.replace('.docx', '')}_Automation_Design.tsu"
        }
    )

@app.get("/download-artifact/execution-report/{file_name:path}")
async def download_execution_report(file_name: str, file_type: str = Query("pdf")):
    """Download execution report (PDF or JSON file)"""
    import urllib.parse
    file_name = urllib.parse.unquote(file_name)
    
    # Use files from Test Execution Agent Output directory
    if file_type.lower() == "json":
        report_file = BASE_DIR / "Test Execution Agent Output" / "P2P E2E Test.json"
        media_type = "application/json"
        filename_suffix = "_Execution_Report.json"
    else:
        report_file = BASE_DIR / "Test Execution Agent Output" / "P2P E2E Test.pdf"
        media_type = "application/pdf"
        filename_suffix = "_Execution_Report.pdf"
    
    print(f"Looking for report file at: {report_file}")
    print(f"File exists: {report_file.exists()}")
    print(f"File type: {file_type}")
    print(f"Absolute path: {report_file.absolute()}")
    
    if not report_file.exists():
        print(f"Error: File not found at {report_file.absolute()}")
        raise HTTPException(status_code=404, detail=f"Execution report not found at {report_file.absolute()}")
    
    # Serve for download
    return FileResponse(
        path=str(report_file),
        filename=f"{file_name.replace('.docx', '')}{filename_suffix}",
        media_type=media_type,
        headers={
            "Content-Disposition": f"attachment; filename={file_name.replace('.docx', '')}{filename_suffix}"
        }
    )

@app.get("/download-outputs/{file_name:path}")
async def download_outputs(file_name: str):
    """Create and download a ZIP file containing all outputs for a file"""
    try:
        # URL decode the file name
        from urllib.parse import unquote
        file_name = unquote(file_name)
        file_stem = Path(file_name).stem
        
        # Create in-memory ZIP file
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add User Story file
            user_story_file = USER_STORY_DIR / f"{file_stem}_userstory.docx"
            if not user_story_file.exists():
                user_story_file = USER_STORY_DIR / f"{file_stem}_userstory.json"
            if user_story_file.exists():
                zip_file.write(user_story_file, f"UserStory/{user_story_file.name}")
            
            # Add Test Cases file
            test_cases_file = TEST_CASES_DIR / f"{file_stem}_testcases.json"
            if test_cases_file.exists():
                zip_file.write(test_cases_file, f"TestCases/{test_cases_file.name}")
            
            # Add Test Data file
            test_data_file = TEST_DATA_DIR / f"{file_stem}_testdata.json"
            if test_data_file.exists():
                zip_file.write(test_data_file, f"TestData/{test_data_file.name}")
            
            # Add Test Coverage file
            coverage_file = TEST_COVERAGE_DIR / f"{file_stem}_testcoverage.json"
            if coverage_file.exists():
                zip_file.write(coverage_file, f"TestCoverage/{coverage_file.name}")
            
            # Add Excel file
            excel_file = EXCEL_OUTPUT_DIR / f"{file_stem}_TestCases_TestData.xlsx"
            if excel_file.exists():
                zip_file.write(excel_file, f"Excel/{excel_file.name}")
        
        zip_buffer.seek(0)
        
        return StreamingResponse(
            io.BytesIO(zip_buffer.read()),
            media_type="application/zip",
            headers={
                "Content-Disposition": f"attachment; filename={file_stem}_ATF_Outputs.zip"
            }
        )
    except Exception as e:
        print(f"Error creating ZIP file: {e}")
        raise HTTPException(status_code=500, detail=f"Error creating ZIP file: {str(e)}")

@app.get("/output-path")
async def get_output_path():
    """Get the output directory path"""
    return {
        "output_path": str(OUTPUT_DIR.absolute()),
        "user_story_path": str(USER_STORY_DIR.absolute()),
        "test_cases_path": str(TEST_CASES_DIR.absolute()),
        "test_data_path": str(TEST_DATA_DIR.absolute()),
        "test_coverage_path": str(TEST_COVERAGE_DIR.absolute()),
        "excel_path": str(EXCEL_OUTPUT_DIR.absolute())
    }

@app.get("/")
async def root():
    return {"message": "ATF (Autonomous Testing Framework) API", "version": "1.0"}

